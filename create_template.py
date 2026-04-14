from docx import Document
from docx.shared import Pt
import os

def create_template():
    doc = Document()
    
    # 1. Title
    title = doc.add_heading('TSCOPE - CONTRATO DE PRESTAÇÃO DE SERVIÇOS', level=0)
    
    # 2. Contractor (Fixed Data)
    p1 = doc.add_paragraph()
    p1.add_run('CONTRATANTE: ').bold = True
    p1.add_run('{contratante_nome}, portador do CPF/CNPJ n. {contratante_doc}, com endereço em {contratante_endereco}.')

    # 3. Contracted (Variable Data)
    p2 = doc.add_paragraph()
    p2.add_run('CONTRATADO: ').bold = True
    p2.add_run('{contratado}, portador do CPF/CNPJ n. {documento_contratado}.')

    # 4. Object
    doc.add_heading('1. DO OBJETO', level=1)
    p3 = doc.add_paragraph()
    p3.add_run('O presente contrato tem como objeto a prestação do seguinte serviço: {objeto_servico}.')

    # 5. Value and Payments
    doc.add_heading('2. DO VALOR E PAGAMENTO', level=1)
    p4 = doc.add_paragraph()
    p4.add_run('Pelo serviço acima descrito, o CONTRATANTE pagará ao CONTRATADO o valor total de ')
    p4.add_run('R$ {valor}').bold = True
    p4.add_run(', nas seguintes condições: {condicoes_pagamento}.')

    # 6. Attachment 1
    doc.add_heading('3. DO ANEXO 1', level=1)
    p5 = doc.add_paragraph()
    p5.add_run('As especificações técnicas detalhadas estão descritas no Anexo 1: ')
    p5.add_run('{anexo_1}').italic = True

    # 7. Signature
    doc.add_paragraph('\n\n\n')
    p6 = doc.add_paragraph()
    p6.add_run('___________________________________________\n')
    p6.add_run('{contratante_nome}\n')
    p6.add_run('CONTRATANTE')

    p7 = doc.add_paragraph('\n\n')
    p7.add_run('___________________________________________\n')
    p7.add_run('{contratado}\n')
    p7.add_run('CONTRATADO')

    filename = 'modelo_teste_tscope.docx'
    doc.save(filename)
    print(f"File created: {filename}")

if __name__ == "__main__":
    create_template()
